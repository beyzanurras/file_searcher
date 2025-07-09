import os
import re
from pathlib import Path
from typing import List, Dict, Tuple

# Dosya okuma kütüphaneleri
try:
    from docx import Document  # python-docx
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


class FileSearcher:
    """Farklı dosya türlerinde anahtar kelime araması yapan sınıf."""
    
    def __init__(self):
        self.supported_extensions = ['.txt', '.docx', '.pdf', '.xlsx']
    
    def search_in_directory(self, directory_path: str, keywords: List[str]) -> List[Dict[str, str]]:
        """
        Belirtilen dizinde anahtar kelimeleri arar.
        
        Args:
            directory_path: Aranacak dizin yolu
            keywords: Aranacak anahtar kelimeler listesi
        
        Returns:
            Bulunan dosyaların bilgilerini içeren liste
        """
        results = []
        
        if not os.path.exists(directory_path):
            return results
        
        # Tüm dosyaları tarar
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = Path(file_path).suffix.lower()
                
                if file_extension in self.supported_extensions:
                    try:
                        content = self._read_file_content(file_path, file_extension)
                        if content and self._search_keywords_in_content(content, keywords):
                            results.append({
                                'file_path': file_path,
                                'file_name': file,
                                'file_type': file_extension,
                                'found_keywords': self._get_found_keywords(content, keywords)
                            })
                    except Exception as e:
                        print(f"Dosya okuma hatası {file_path}: {str(e)}")
        
        return results
    
    def _read_file_content(self, file_path: str, file_extension: str) -> str:
        """
        Dosya türüne göre içeriği okur.
        
        Args:
            file_path: Dosya yolu
            file_extension: Dosya uzantısı
        
        Returns:
            Dosya içeriği (string)
        """
        content = ""
        
        try:
            if file_extension == '.txt':
                content = self._read_txt_file(file_path)
            elif file_extension == '.docx':
                content = self._read_docx_file(file_path)
            elif file_extension == '.pdf':
                content = self._read_pdf_file(file_path)
            elif file_extension == '.xlsx':
                content = self._read_xlsx_file(file_path)
        except Exception as e:
            print(f"Dosya okuma hatası {file_path}: {str(e)}")
            return ""
        
        return content
    
    def _read_txt_file(self, file_path: str) -> str:
        """TXT dosyasını okur."""
        encodings = ['utf-8', 'utf-8-sig', 'iso-8859-9', 'windows-1254']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
            except Exception:
                break
        
        return ""
    
    def _read_docx_file(self, file_path: str) -> str:
        """DOCX dosyasını okur."""
        if Document is None:
            return ""
        
        try:
            doc = Document(file_path)
            content = []
            
            # Paragrafları okur
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            
            # Tabloları okur
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content.append(cell.text)
            
            return '\n'.join(content)
        except Exception:
            return ""
    
    def _read_pdf_file(self, file_path: str) -> str:
        """PDF dosyasını okur."""
        if PyPDF2 is None:
            return ""
        
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    try:
                        text = page.extract_text()
                        if text:
                            content.append(text)
                    except Exception:
                        continue
            
            return '\n'.join(content)
        except Exception:
            return ""
    
    def _read_xlsx_file(self, file_path: str) -> str:
        """XLSX dosyasını okur."""
        if load_workbook is None:
            return ""
        
        try:
            workbook = load_workbook(file_path, data_only=True)
            content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                for row in sheet.iter_rows(values_only=True):
                    for cell_value in row:
                        if cell_value is not None:
                            content.append(str(cell_value))
            
            return '\n'.join(content)
        except Exception:
            return ""
    
    def _search_keywords_in_content(self, content: str, keywords: List[str]) -> bool:
        """
        İçerikte anahtar kelimeleri arar.
        
        Args:
            content: Aranacak metin
            keywords: Anahtar kelimeler
        
        Returns:
            Herhangi bir anahtar kelime bulundu mu
        """
        if not content or not keywords:
            return False
        
        content_lower = content.lower()
        
        for keyword in keywords:
            if keyword.lower().strip() in content_lower:
                return True
        
        return False
    
    def _get_found_keywords(self, content: str, keywords: List[str]) -> List[str]:
        """
        İçerikte bulunan anahtar kelimeleri döndürür.
        
        Args:
            content: Aranacak metin
            keywords: Anahtar kelimeler
        
        Returns:
            Bulunan anahtar kelimeler listesi
        """
        found_keywords = []
        
        if not content or not keywords:
            return found_keywords
        
        content_lower = content.lower()
        
        for keyword in keywords:
            if keyword.lower().strip() in content_lower:
                found_keywords.append(keyword.strip())
        
        return found_keywords 
